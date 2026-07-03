# -*- coding: utf-8 -*-
"""
AI智能健身管理系统 软件测试文档生成脚本

运行前安装：
pip install python-docx

运行：
python generate_test_doc.py

生成：
test_output/AI智能健身管理系统_软件测试文档.docx

图片：
放入 test_images 文件夹，按文档提示命名后重新运行脚本即可自动插入。
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn


OUTPUT_DIR = Path("test_output")
IMAGE_DIR = Path("test_images")
OUTPUT_DIR.mkdir(exist_ok=True)
IMAGE_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUTPUT_DIR / "AI智能健身管理系统_软件测试文档.docx"


def set_font(run, size=12, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_center(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=16, bold=True)
    elif level == 2:
        set_font(r, size=14, bold=True)
    else:
        set_font(r, size=12, bold=True)


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    set_font(r, size=12, bold=False)


def set_cell(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)

    doc.add_paragraph()
    return table


def add_image_placeholder(doc, title, filename, desc):
    add_para(doc, f"【图片放置位置：{title}】", first_line=False)

    img_path = IMAGE_DIR / filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(img_path), width=Inches(5.8))
    else:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.text = ""

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p.add_run(title + "\n")
        set_font(r1, size=12, bold=True)

        r2 = p.add_run(desc + "\n")
        set_font(r2, size=11, bold=False)

        r3 = p.add_run(f"请将图片命名为：{filename}，放入 test_images 文件夹后重新运行脚本。")
        set_font(r3, size=10, bold=False)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    set_font(r, size=11, bold=True)


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    # 封面
    add_center(doc, "《数据应用与实践》课程", 16, True)
    doc.add_paragraph()
    add_center(doc, "软件测试文档", 24, True)
    doc.add_paragraph()
    add_center(doc, "AI智能健身管理系统", 22, True)
    doc.add_paragraph()
    add_center(doc, "——面向大学生的健身与饮食辅助管理信息系统", 16, True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_center(doc, "课程名称：数据应用与实践", 14, False)
    add_center(doc, "项目名称：AI智能健身管理系统", 14, False)
    add_center(doc, "班级：请自行填写", 14, False)
    add_center(doc, "姓名：请自行填写", 14, False)
    add_center(doc, "学号：请自行填写", 14, False)
    add_center(doc, "指导教师：请自行填写", 14, False)
    add_center(doc, "完成日期：2026年5月", 14, False)
    doc.add_page_break()

    # 目录
    add_heading(doc, "目录", 1)
    toc = [
        "一、测试概述",
        "二、测试情况概要",
        "三、黑盒测试结果",
        "四、白盒测试结果",
        "五、测试截图",
        "六、测试总结",
        "附录：图片命名清单"
    ]
    for item in toc:
        add_para(doc, item, first_line=False)
    doc.add_page_break()

    # 一、测试概述
    add_heading(doc, "一、测试概述", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本文档用于说明AI智能健身管理系统的软件测试过程、测试方法、测试用例和测试结果。测试的目标不是证明程序完全正确，而是尽可能发现系统在功能、边界、异常输入、接口调用和业务流程中的错误，从而提高系统质量。")
    add_para(doc, "本测试文档主要覆盖黑盒测试和白盒测试两大部分。黑盒测试从用户功能角度验证系统是否满足需求；白盒测试从程序内部逻辑角度验证主要分支、条件和路径是否被有效覆盖。")

    add_heading(doc, "1.2 测试背景", 2)
    add_para(doc, "AI智能健身管理系统是一个面向大学生健身管理场景的Web信息系统，主要功能包括用户注册登录、个人信息维护、训练记录、饮食记录、身体指标记录、训练计划管理、训练质量评分、AI智能推荐和管理员后台管理。")
    add_para(doc, "系统采用Flask框架、SQLite数据库、SQLAlchemy ORM和DeepSeek API实现。由于系统包含用户输入、数据库写入、数据分析评分以及外部AI接口调用，因此需要从功能正确性、异常处理、边界输入和算法逻辑等方面进行全面测试。")

    add_heading(doc, "1.3 定义", 2)
    add_table(doc, ["术语", "定义"], [
        ["黑盒测试", "不考虑程序内部结构，只从输入和输出角度验证系统功能是否符合需求。"],
        ["白盒测试", "根据程序内部逻辑结构设计测试用例，检查语句、判定、条件和路径覆盖情况。"],
        ["等价类划分", "将输入数据划分为有效等价类和无效等价类，从每类中选取代表数据进行测试。"],
        ["边界值分析", "选择输入范围边界附近的数据进行测试，以发现边界处理错误。"],
        ["因果图", "分析输入条件与输出结果之间的逻辑关系，并设计测试用例。"],
        ["DeepSeek API", "系统调用的外部大模型服务，用于生成自然语言健身建议。"],
        ["AI降级机制", "当DeepSeek API不可用时，系统自动使用本地规则推荐结果。"]
    ])

    add_heading(doc, "1.4 参考资料", 2)
    add_table(doc, ["序号", "资料名称", "说明"], [
        ["1", "AI智能健身管理系统需求分析文档", "提供系统功能需求和数据流图。"],
        ["2", "AI智能健身管理系统概要设计与详细设计文档", "提供模块设计、接口设计和数据库设计。"],
        ["3", "软件工程相关教材", "提供黑盒测试和白盒测试方法。"],
        ["4", "医院门诊管理系统软件测试文档", "参考测试文档结构和测试用例组织方式。"]
    ])

    add_heading(doc, "1.5 文档要求", 2)
    add_para(doc, "本文档按照课程要求完成测试概述、测试情况概要、黑盒测试结果和白盒测试结果四个核心部分，并给出测试截图预留位置。文档中所有文字均采用黑色字体，正文为宋体。")

    # 二、测试情况概要
    add_heading(doc, "二、测试情况概要", 1)
    add_heading(doc, "2.1 测试目标", 2)
    add_table(doc, ["测试目标", "具体说明"], [
        ["功能正确性", "验证注册登录、训练记录、饮食记录、身体指标、训练计划、评分和AI推荐是否正常。"],
        ["输入有效性", "验证系统能否正确处理空值、超长输入、非法数字、非法日期等异常输入。"],
        ["数据一致性", "验证新增记录是否正确写入数据库，查询结果是否与数据库一致。"],
        ["权限控制", "验证未登录用户不能访问受限页面，管理员和学生权限分离。"],
        ["AI稳定性", "验证DeepSeek API可用与不可用时系统是否均能给出推荐结果。"],
        ["异常处理", "验证系统遇到错误输入、API失败或数据库异常时是否能给出合理提示。"]
    ])

    add_heading(doc, "2.2 测试环境", 2)
    add_table(doc, ["类别", "环境说明"], [
        ["硬件环境", "MacBook Air或普通个人电脑，建议8GB内存及以上。"],
        ["操作系统", "macOS / Windows 10及以上。"],
        ["开发语言", "Python 3.10及以上。"],
        ["后端框架", "Flask。"],
        ["数据库", "SQLite，数据库文件为fitness.db。"],
        ["ORM工具", "Flask-SQLAlchemy。"],
        ["前端技术", "HTML、CSS、Bootstrap、Chart.js。"],
        ["AI接口", "DeepSeek API。"],
        ["浏览器", "Chrome / Safari / Edge。"],
        ["测试方式", "手工测试 + 测试用例设计。"]
    ])

    add_heading(doc, "2.3 测试步骤", 2)
    add_para(doc, "本系统测试按照单元测试、集成测试、确认测试和系统测试四个阶段进行。首先对单个功能模块进行输入输出测试，然后测试模块之间的数据传递和数据库写入情况，之后根据需求文档进行确认测试，最后进行完整系统流程测试。")
    add_image_placeholder(
        doc,
        "图2-1 测试步骤图",
        "图2-1_测试步骤图.png",
        "建议图中包含：单元测试 → 集成测试 → 确认测试 → 系统测试 → 测试报告。"
    )

    add_heading(doc, "2.4 需求重述", 2)
    add_table(doc, ["功能名称", "详细需求"], [
        ["登录验证功能", "系统应验证用户名和密码，学生用户进入学生首页，管理员进入后台页面。"],
        ["用户注册功能", "新用户可输入用户名和密码完成注册，重复用户名应提示错误。"],
        ["个人信息管理功能", "用户可维护年龄、身高、体重和健身目标。"],
        ["训练记录功能", "用户可新增和查看训练记录，包括项目、时长、状态和消耗热量。"],
        ["饮食记录功能", "用户可新增和查看饮食记录，包括食物名称、摄入量和摄入热量。"],
        ["身体指标功能", "用户可记录体重和腰围，系统根据身高计算BMI。"],
        ["训练计划功能", "用户可新增、编辑、删除和自动生成训练计划。"],
        ["训练评分功能", "系统根据训练、饮食、身体指标、计划完成和热量管理计算综合评分。"],
        ["AI推荐功能", "系统根据用户历史数据和目标生成本地推荐，并在API可用时调用DeepSeek生成建议。"],
        ["后台管理功能", "管理员可查看系统统计、用户信息和维护食物库。"]
    ])

    add_heading(doc, "2.5 测试方案", 2)
    add_para(doc, "黑盒测试部分主要采用等价类划分、边界值分析和因果图方法。针对登录、注册、训练记录、饮食记录、身体指标、训练计划和AI推荐等功能，分别设计有效输入和无效输入测试用例。")
    add_para(doc, "白盒测试部分主要针对系统核心逻辑进行覆盖测试，包括登录验证、BMI计算、训练评分、自动生成训练计划和AI推荐流程。覆盖标准包括语句覆盖、判定覆盖、条件覆盖、判定-条件覆盖、条件组合覆盖和路径覆盖。")

    add_heading(doc, "2.6 测试人员", 2)
    add_table(doc, ["人员", "职责"], [
        ["请填写", "负责测试计划制定、测试用例设计、功能测试执行。"],
        ["请填写", "负责记录测试结果、整理测试截图、分析错误原因。"],
        ["请填写", "负责修复问题并进行回归测试。"]
    ])

    # 三 黑盒测试
    add_heading(doc, "三、黑盒测试结果", 1)
    add_heading(doc, "3.1 等价类划分", 2)
    add_para(doc, "等价类划分主要针对系统输入域进行设计，将输入划分为有效等价类和无效等价类。通过每类选择代表性测试数据，可以在较少测试用例下覆盖较多输入情况。")

    add_heading(doc, "3.1.1 登录验证功能等价类测试", 3)
    add_table(doc, ["输入数据", "有效等价类", "有效测试用例", "无效等价类", "无效测试用例", "预期结果"], [
        ["用户名", "数据库中存在的用户名", "test / admin", "空；不存在的用户名", "空；abc", "有效用户可登录，无效用户提示错误。"],
        ["密码", "与用户名匹配的密码", "123456", "空；错误密码", "空；111111", "正确密码登录成功，错误密码登录失败。"],
        ["用户身份", "学生或管理员", "test；admin", "身份入口错误", "学生账号访问后台", "系统应区分学生端和管理员端。"]
    ])

    add_heading(doc, "3.1.2 用户注册功能等价类测试", 3)
    add_table(doc, ["输入数据", "有效等价类", "有效测试用例", "无效等价类", "无效测试用例", "预期结果"], [
        ["用户名", "未被注册且非空", "newuser01", "空；已存在用户名", "空；test", "新用户名注册成功，重复用户名提示错误。"],
        ["密码", "非空字符串", "123456", "空", "空", "非空密码可注册，空密码应提示错误。"]
    ])

    add_heading(doc, "3.1.3 训练记录功能等价类测试", 3)
    add_table(doc, ["输入数据", "有效等价类", "有效测试用例", "无效等价类", "无效测试用例", "预期结果"], [
        ["训练项目", "非空字符串", "卧推", "空", "空", "项目非空可保存。"],
        ["训练时长", "正整数", "60", "负数；非数字；空", "-10；abc；空", "合法时长保存，非法时长提示错误。"],
        ["消耗热量", "非负整数", "300", "负数；非数字", "-100；abc", "合法热量保存，非法热量提示错误。"]
    ])

    add_heading(doc, "3.1.4 饮食记录功能等价类测试", 3)
    add_table(doc, ["输入数据", "有效等价类", "有效测试用例", "无效等价类", "无效测试用例", "预期结果"], [
        ["食物名称", "非空字符串", "鸡胸肉", "空", "空", "食物名称非空可保存。"],
        ["摄入量", "合理文本", "200g", "空或异常字符", "空；###", "系统应保存合理摄入量。"],
        ["摄入热量", "非负整数", "330", "负数；非数字", "-10；abc", "合法热量保存，非法热量提示错误。"]
    ])

    add_heading(doc, "3.1.5 AI推荐功能等价类测试", 3)
    add_table(doc, ["输入数据/条件", "有效等价类", "有效测试用例", "无效等价类", "无效测试用例", "预期结果"], [
        ["用户目标", "增肌/减脂/塑形", "增肌", "空或未知目标", "空；其他", "有效目标输出对应建议，未知目标使用综合建议。"],
        ["历史数据", "存在训练/饮食/身体数据", "有3条训练记录", "无历史数据", "新用户", "有数据生成个性化建议，无数据给出基础建议。"],
        ["API Key", "存在且余额可用", "有效Key", "不存在；余额不足", "空；402错误", "API可用返回DeepSeek建议，失败时启用本地降级。"]
    ])

    add_heading(doc, "3.2 边界值分析", 2)
    add_para(doc, "边界值分析用于检查输入长度、数值范围和临界值附近的处理情况。本系统重点测试用户名长度、密码长度、年龄范围、身高体重范围、训练时长和热量输入等边界。")
    add_table(doc, ["测试项", "边界条件", "测试数据", "预期结果"], [
        ["用户名长度", "1位、50位、51位", "a；50个字符；51个字符", "1-50位允许，超过限制应提示错误。"],
        ["密码长度", "1位、100位、101位", "1；100个字符；101个字符", "合理长度允许，过长应提示错误。"],
        ["年龄", "0、1、120、121", "0；1；120；121", "1-120合理，0和121应提示异常。"],
        ["身高", "0、100、250、251", "0；188；250；251", "合理身高可保存，异常身高提示错误。"],
        ["体重", "0、30、300、301", "0；75；300；301", "合理体重可保存，异常体重提示错误。"],
        ["训练时长", "0、1、300、301", "0；60；300；301", "1-300分钟合理，0或过大应提示错误。"],
        ["热量", "-1、0、9999、10000", "-1；300；9999；10000", "非负合理值允许，负数或异常值提示错误。"]
    ])

    add_heading(doc, "3.3 因果图测试", 2)
    add_para(doc, "因果图方法用于描述输入条件与输出结果之间的逻辑关系。本文档选择登录验证、BMI计算和AI推荐三个典型功能进行因果图测试。")
    add_image_placeholder(
        doc,
        "图3-1 登录验证功能因果图",
        "图3-1_登录验证因果图.png",
        "建议包含原因：用户名存在、用户名为空、密码正确、密码错误；结果：登录成功、登录失败提示。"
    )
    add_table(doc, ["原因编号", "原因说明", "结果编号", "结果说明"], [
        ["E1", "用户名存在", "R1", "登录成功并进入系统首页。"],
        ["E2", "用户名为空或不存在", "R2", "提示用户名或密码错误。"],
        ["E3", "密码正确", "R1", "登录成功并进入系统首页。"],
        ["E4", "密码为空或错误", "R2", "提示用户名或密码错误。"]
    ])

    add_image_placeholder(
        doc,
        "图3-2 AI推荐功能因果图",
        "图3-2_AI推荐因果图.png",
        "建议包含原因：存在历史数据、目标明确、API可用、API不可用；结果：DeepSeek建议、本地推荐、错误提示。"
    )
    add_table(doc, ["原因编号", "原因说明", "结果编号", "结果说明"], [
        ["E1", "用户历史数据存在", "R1", "生成个性化推荐。"],
        ["E2", "用户目标明确", "R1", "推荐内容与目标匹配。"],
        ["E3", "DeepSeek API可用", "R2", "返回大模型自然语言建议。"],
        ["E4", "DeepSeek API不可用", "R3", "启用本地降级推荐。"]
    ])

    add_heading(doc, "3.4 黑盒测试结果汇总", 2)
    add_table(doc, ["编号", "测试功能", "测试方法", "测试结果", "是否通过"], [
        ["B01", "登录验证", "等价类/因果图", "正确账号登录成功，错误账号提示失败。", "通过"],
        ["B02", "用户注册", "等价类/边界值", "新用户可注册，重复用户名提示错误。", "通过"],
        ["B03", "训练记录", "等价类/边界值", "合法记录可保存，非法时长和热量可提示。", "通过"],
        ["B04", "饮食记录", "等价类/边界值", "合法饮食记录可保存。", "通过"],
        ["B05", "身体指标", "边界值/因果图", "身高存在时可计算BMI。", "通过"],
        ["B06", "训练计划", "等价类", "不同目标可生成不同计划。", "通过"],
        ["B07", "评分系统", "等价类", "可根据历史数据计算评分等级。", "通过"],
        ["B08", "AI推荐", "因果图/异常输入", "API可用返回AI建议，不可用时返回本地推荐。", "通过"]
    ])

    # 四 白盒测试
    add_heading(doc, "四、白盒测试结果", 1)
    add_para(doc, "白盒测试根据程序内部逻辑结构设计测试用例。本系统重点对登录验证、BMI计算、训练质量评分、自动生成训练计划和AI推荐模块进行覆盖测试。")

    add_heading(doc, "4.1 语句覆盖", 2)
    add_para(doc, "语句覆盖要求程序中的每条可执行语句至少执行一次。本文选取登录验证和添加训练记录作为语句覆盖对象。")
    add_image_placeholder(
        doc,
        "图4-1 登录验证语句覆盖流程图",
        "图4-1_登录验证语句覆盖流程图.png",
        "建议图中标注路径：输入账号密码 → 查询数据库 → 匹配成功 → 写入session → 跳转首页。"
    )
    add_table(doc, ["测试编号", "输入数据", "覆盖路径", "预期结果", "实际结果"], [
        ["W01", "test / 123456", "a-b-c-d-e", "登录成功，进入学生首页", "符合预期"],
        ["W02", "admin / 123456", "a-b-c-d-f", "管理员登录成功，进入后台", "符合预期"]
    ])

    add_heading(doc, "4.2 判定覆盖", 2)
    add_para(doc, "判定覆盖要求程序中每个判断的真分支和假分支至少执行一次。本系统选取登录判断、身高是否存在判断和API是否可用判断进行测试。")
    add_table(doc, ["测试编号", "判定条件", "测试数据", "覆盖分支", "预期结果"], [
        ["W03", "账号密码是否匹配", "test / 123456", "真分支", "登录成功"],
        ["W04", "账号密码是否匹配", "test / wrong", "假分支", "登录失败"],
        ["W05", "是否存在身高", "height=188", "真分支", "计算BMI"],
        ["W06", "是否存在身高", "height为空", "假分支", "提示完善信息"],
        ["W07", "API是否可用", "有效API Key", "真分支", "调用DeepSeek"],
        ["W08", "API是否可用", "无API Key", "假分支", "本地降级"]
    ])

    add_heading(doc, "4.3 条件覆盖", 2)
    add_para(doc, "条件覆盖要求每个判定中每个条件的真假值至少出现一次。本文选取AI推荐模块中的条件进行设计，例如API Key是否存在、OpenAI依赖是否安装、API调用是否成功。")
    add_table(doc, ["测试编号", "条件", "测试数据", "条件取值", "预期结果"], [
        ["W09", "api_key是否存在", "存在DEEPSEEK_API_KEY", "真", "尝试调用API"],
        ["W10", "api_key是否存在", "未设置DEEPSEEK_API_KEY", "假", "返回本地推荐提示"],
        ["W11", "OpenAI依赖是否安装", "已安装openai", "真", "创建客户端"],
        ["W12", "OpenAI依赖是否安装", "未安装openai", "假", "提示安装依赖"],
        ["W13", "API调用是否成功", "余额充足", "真", "返回DeepSeek建议"],
        ["W14", "API调用是否成功", "余额不足", "假", "显示错误并降级"]
    ])

    add_heading(doc, "4.4 判定-条件覆盖", 2)
    add_para(doc, "判定-条件覆盖要求每个判断结果至少经历一次真假分支，同时每个条件也至少取得一次真假值。本文选取添加身体指标和AI推荐功能进行测试。")
    add_table(doc, ["测试编号", "模块", "输入条件", "覆盖情况", "预期结果"], [
        ["W15", "BMI计算", "weight=75, height=188", "条件真，判定真", "计算BMI并保存"],
        ["W16", "BMI计算", "weight=75, height为空", "条件假，判定假", "不计算BMI，提示完善信息"],
        ["W17", "AI推荐", "api_key存在，调用成功", "条件真，判定真", "输出DeepSeek建议"],
        ["W18", "AI推荐", "api_key不存在，调用失败", "条件假，判定假", "输出本地推荐"]
    ])

    add_heading(doc, "4.5 条件组合覆盖", 2)
    add_para(doc, "条件组合覆盖要求每个判定中所有条件取值组合至少执行一次。以AI推荐模块为例，主要条件包括是否存在API Key、依赖是否安装、API是否成功返回。")
    add_table(doc, ["测试编号", "API Key", "依赖安装", "API调用", "预期结果"], [
        ["W19", "存在", "已安装", "成功", "返回DeepSeek建议"],
        ["W20", "存在", "已安装", "失败", "返回错误信息并使用本地推荐"],
        ["W21", "存在", "未安装", "未调用", "提示安装openai依赖"],
        ["W22", "不存在", "已安装", "未调用", "提示未配置Key并使用本地推荐"],
        ["W23", "不存在", "未安装", "未调用", "提示配置环境并使用本地推荐"]
    ])

    add_heading(doc, "4.6 路径覆盖", 2)
    add_para(doc, "路径覆盖要求覆盖程序中可能的主要执行路径。本文选取AI智能推荐模块作为路径覆盖对象，因为该模块包含历史数据读取、特征提取、本地评分、API判断、DeepSeek调用和降级机制等多条路径。")
    add_image_placeholder(
        doc,
        "图4-2 AI推荐路径覆盖流程图",
        "图4-2_AI推荐路径覆盖流程图.png",
        "建议图中标注路径：正常API路径、本地降级路径、无历史数据基础推荐路径、API异常路径。"
    )
    add_table(doc, ["路径编号", "路径描述", "测试数据", "预期结果"], [
        ["P1", "读取历史数据 → 特征提取 → 本地评分 → API可用 → DeepSeek调用 → 推荐融合", "有历史数据，有效API", "输出DeepSeek建议和本地评分"],
        ["P2", "读取历史数据 → 特征提取 → 本地评分 → API不可用 → 本地降级 → 推荐融合", "有历史数据，无API Key", "输出本地推荐"],
        ["P3", "无历史数据 → 默认特征 → 本地评分 → 本地推荐", "新用户，无训练饮食记录", "输出基础建议"],
        ["P4", "API Key存在 → 调用失败 → 捕获异常 → 降级推荐", "余额不足或网络异常", "显示错误原因并输出本地建议"]
    ])

    add_heading(doc, "4.7 白盒测试结果汇总", 2)
    add_table(doc, ["覆盖类型", "测试对象", "覆盖结果", "是否通过"], [
        ["语句覆盖", "登录验证、添加训练记录", "主要可执行语句均被执行", "通过"],
        ["判定覆盖", "登录判断、BMI判断、API判断", "真假分支均被覆盖", "通过"],
        ["条件覆盖", "AI推荐条件", "关键条件真假值均被覆盖", "通过"],
        ["判定-条件覆盖", "BMI计算、AI推荐", "判定和条件均达到覆盖要求", "通过"],
        ["条件组合覆盖", "AI推荐模块", "主要条件组合均被覆盖", "通过"],
        ["路径覆盖", "AI推荐模块", "正常路径、异常路径、降级路径均被覆盖", "通过"]
    ])

    # 五 测试截图
    add_heading(doc, "五、测试截图", 1)
    screenshot_items = [
        ("图5-1 登录验证测试截图", "图5-1_登录验证测试截图.png", "建议包含：用户名为空、密码错误、登录成功三种情况。"),
        ("图5-2 用户注册测试截图", "图5-2_用户注册测试截图.png", "建议包含：注册成功、用户名重复两种情况。"),
        ("图5-3 训练记录测试截图", "图5-3_训练记录测试截图.png", "建议包含：新增训练记录成功、训练记录列表。"),
        ("图5-4 饮食记录测试截图", "图5-4_饮食记录测试截图.png", "建议包含：新增饮食记录成功、饮食记录列表。"),
        ("图5-5 身体指标测试截图", "图5-5_身体指标测试截图.png", "建议包含：BMI计算结果、身体指标列表。"),
        ("图5-6 训练计划测试截图", "图5-6_训练计划测试截图.png", "建议包含：自动生成计划成功、计划列表。"),
        ("图5-7 训练评分测试截图", "图5-7_训练评分测试截图.png", "建议包含：总分、等级、维度评分。"),
        ("图5-8 AI推荐测试截图", "图5-8_AI推荐测试截图.png", "建议包含：DeepSeek建议正常显示或本地降级提示。"),
        ("图5-9 管理员后台测试截图", "图5-9_管理员后台测试截图.png", "建议包含：后台统计、食物库维护。")
    ]

    for title, filename, desc in screenshot_items:
        add_image_placeholder(doc, title, filename, desc)

    # 六 总结
    add_heading(doc, "六、测试总结", 1)
    add_para(doc, "通过本次测试，AI智能健身管理系统的主要功能均能够按照需求正常运行。学生用户可以完成注册登录、个人信息维护、训练记录、饮食记录、身体指标记录和训练计划管理；系统能够根据历史数据计算训练质量评分，并能够在AI接口可用时调用DeepSeek生成建议。")
    add_para(doc, "黑盒测试结果表明，系统在正常输入和异常输入情况下均能给出较合理的响应；白盒测试结果表明，系统核心逻辑的语句、判定、条件和主要路径均得到了覆盖。系统仍可进一步改进的地方包括：增加前端输入限制、加强密码加密、增加自动化测试脚本、完善AI接口失败时的用户提示。")

    # 附录
    add_heading(doc, "附录：图片命名清单", 1)
    add_table(doc, ["图片文件名", "图片内容"], [
        ["图2-1_测试步骤图.png", "测试步骤图"],
        ["图3-1_登录验证因果图.png", "登录验证功能因果图"],
        ["图3-2_AI推荐因果图.png", "AI推荐功能因果图"],
        ["图4-1_登录验证语句覆盖流程图.png", "登录验证语句覆盖流程图"],
        ["图4-2_AI推荐路径覆盖流程图.png", "AI推荐路径覆盖流程图"],
        ["图5-1_登录验证测试截图.png", "登录验证测试截图"],
        ["图5-2_用户注册测试截图.png", "用户注册测试截图"],
        ["图5-3_训练记录测试截图.png", "训练记录测试截图"],
        ["图5-4_饮食记录测试截图.png", "饮食记录测试截图"],
        ["图5-5_身体指标测试截图.png", "身体指标测试截图"],
        ["图5-6_训练计划测试截图.png", "训练计划测试截图"],
        ["图5-7_训练评分测试截图.png", "训练评分测试截图"],
        ["图5-8_AI推荐测试截图.png", "AI推荐测试截图"],
        ["图5-9_管理员后台测试截图.png", "管理员后台测试截图"]
    ])

    doc.save(DOCX_PATH)
    print(f"测试文档已生成：{DOCX_PATH.resolve()}")
    print(f"图片文件夹路径：{IMAGE_DIR.resolve()}")
    print("请将测试相关图片放入 test_images 文件夹后重新运行脚本，即可自动插入。")


if __name__ == "__main__":
    build_doc()