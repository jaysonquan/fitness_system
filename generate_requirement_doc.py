# -*- coding: utf-8 -*-
"""
AI智能健身管理系统需求分析文档生成脚本 V3
特点：
1. 内容详细全面，参考结构化需求分析文档格式
2. 标题全部黑体，正文全部宋体
3. 所有文字颜色均为黑色
4. 自动生成图片放置位置标注
5. 如果你手动画好图片并放入 images 文件夹，脚本会自动插入
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_DIR = Path("requirement_output_v3")
IMAGE_DIR = Path("images")
OUTPUT_DIR.mkdir(exist_ok=True)
IMAGE_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUTPUT_DIR / "AI智能健身管理系统需求分析文档_详细版.docx"


def set_run_font(run, font_name="宋体", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(paragraph, first_line=True, center=False):
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(6)

    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def add_cover_line(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, "黑体" if bold else "宋体", size, bold)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    r = p.add_run(text)
    size = 16 if level == 1 else 14
    set_run_font(r, "黑体", size, True)


def add_paragraph(doc, text, first_line=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    r = p.add_run(text)
    set_run_font(r, "宋体", 12, False)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run("（1）" + text)
    set_run_font(r, "宋体", 12, False)


def add_numbered_item(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(f"{number}. {text}")
    set_run_font(r, "宋体", 12, False)


def shade_cell(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_run_font(r, "宋体", 10, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False)

    doc.add_paragraph()
    return table


def add_image_placeholder(doc, title, filename, description):
    add_paragraph(doc, f"【图片放置位置：{title}】", first_line=False)

    image_path = IMAGE_DIR / filename

    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.8))
    else:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p.add_run(title + "\n")
        set_run_font(r1, "黑体", 12, True)

        r2 = p.add_run(description + "\n")
        set_run_font(r2, "宋体", 11, False)

        r3 = p.add_run(f"请将手绘或draw.io导出的图片命名为：{filename}，放入 images 文件夹后重新运行脚本。")
        set_run_font(r3, "宋体", 10, False)

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run(title)
    set_run_font(r, "黑体", 11, True)


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    # 封面
    add_cover_line(doc, "《数据应用与实践》课程", 16, True)
    doc.add_paragraph()
    add_cover_line(doc, "需求分析文档", 22, True)
    doc.add_paragraph()
    add_cover_line(doc, "AI智能健身管理系统", 24, True)
    doc.add_paragraph()
    add_cover_line(doc, "——面向大学生的健身与饮食辅助管理信息系统", 16, True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_cover_line(doc, "课程名称：数据应用与实践", 14)
    add_cover_line(doc, "项目名称：AI智能健身管理系统", 14)
    add_cover_line(doc, "班级：请自行填写", 14)
    add_cover_line(doc, "姓名：请自行填写", 14)
    add_cover_line(doc, "学号：请自行填写", 14)
    add_cover_line(doc, "指导教师：请自行填写", 14)
    add_cover_line(doc, "完成日期：2026年5月", 14)
    doc.add_page_break()

    # 目录
    add_heading(doc, "目录", 1)
    toc_items = [
        "一、项目人员信息",
        "二、需求获取",
        "三、系统用户说明",
        "四、可行性分析",
        "五、开发与运行环境",
        "六、需求分析",
        "七、系统功能需求表",
        "八、系统功能建模：分层数据流图",
        "九、数据字典",
        "十、系统数据建模：E-R模型",
        "十一、关系模式",
        "十二、非功能需求",
        "十三、创新点总结",
        "十四、项目组成员分工"
    ]
    for item in toc_items:
        add_paragraph(doc, item, first_line=False)
    doc.add_page_break()

    # 一
    add_heading(doc, "一、项目人员信息", 1)
    add_table(doc, ["姓名", "学号", "班级", "项目分工"], [
        ["请填写", "请填写", "请填写", "后端开发、数据库设计、AI推荐模块"],
        ["请填写", "请填写", "请填写", "前端页面、数据可视化、文档整理"],
    ])

    # 二
    add_heading(doc, "二、需求获取", 1)
    add_heading(doc, "2.1 项目背景", 2)
    add_paragraph(doc, "随着健康生活理念的普及，越来越多大学生开始关注健身训练、体重管理与科学饮食。然而，在实际生活中，许多学生的健身行为仍然依赖短视频、网络帖子或个人经验，缺少系统化、持续化和个性化的指导。部分学生虽然能够坚持训练，但缺乏对训练频率、饮食热量、身体指标变化和计划完成情况的综合分析，因此难以及时判断训练效果。")
    add_paragraph(doc, "传统信息管理系统通常只能完成数据录入、查询和维护，主要解决“数据存储”的问题，而无法进一步提供智能分析和决策支持。针对这一不足，本项目设计AI智能健身管理系统，在传统健身信息管理功能基础上，引入数据分析模型、训练质量评分模型和DeepSeek大模型接口，实现从“数据记录”到“数据分析”，再到“智能推荐”的升级。")
    add_paragraph(doc, "本系统面向大学生群体，结合校园健身、体育课程训练管理、学生健康行为记录等真实应用场景，提供用户管理、训练记录、饮食记录、身体指标记录、训练计划、数据可视化和AI个性化建议等功能，具有较强现实意义和应用价值。")

    add_heading(doc, "2.2 项目目标", 2)
    goals = [
        "为大学生提供统一的健身与饮食数据管理平台。",
        "支持学生记录训练、饮食和身体指标，实现个人健康数据持续积累。",
        "通过图表展示训练频率、热量收支、体重变化和计划完成率。",
        "构建训练质量评分模型，对用户训练状态进行量化评价。",
        "接入DeepSeek大模型，根据用户数据生成个性化训练与饮食建议。",
        "设计AI降级机制，在大模型接口不可用时仍能通过本地模型提供推荐结果。",
        "为管理员提供用户管理、食物库管理和系统数据维护功能。"
    ]
    for i, item in enumerate(goals, 1):
        add_numbered_item(doc, i, item)

    add_heading(doc, "2.3 项目意义", 2)
    add_paragraph(doc, "第一，本系统有助于提升大学生健康管理的科学性。系统通过持续记录训练、饮食和身体指标，将用户健身行为由碎片化经验管理转变为数据驱动管理。")
    add_paragraph(doc, "第二，本系统有助于提高健身管理效率。用户可以在一个平台中完成训练计划查看、训练记录录入、饮食热量记录和身体指标跟踪，避免数据分散。")
    add_paragraph(doc, "第三，本系统体现了数据应用与人工智能技术在真实场景中的结合。系统不仅完成基础CRUD功能，还通过评分模型、大模型推荐和可视化分析增强系统智能化水平。")
    add_paragraph(doc, "第四，本系统具备良好扩展性。后续可扩展运动动作库、营养计算模型、校园体育课程管理、可穿戴设备数据接入和云端部署。")

    # 三
    add_heading(doc, "三、系统用户说明", 1)
    add_heading(doc, "3.1 学生用户", 2)
    student_needs = [
        "注册和登录系统，建立个人账号。",
        "维护个人基础信息，包括年龄、身高、体重和健身目标。",
        "记录每日训练情况，包括训练项目、训练时长、完成状态和消耗热量。",
        "记录饮食信息，包括食物名称、摄入量和热量。",
        "记录身体指标，包括体重、BMI和腰围。",
        "查看系统生成的训练计划和智能推荐建议。",
        "查看训练评分、热量分析、体重趋势和计划完成率。"
    ]
    for i, item in enumerate(student_needs, 1):
        add_numbered_item(doc, i, item)

    add_heading(doc, "3.2 管理员", 2)
    admin_needs = [
        "登录后台管理系统。",
        "查看和管理学生用户信息。",
        "维护食物库及热量数据。",
        "查看系统整体统计信息。",
        "管理系统基础数据，保证数据完整性。",
        "对系统异常数据进行检查和维护。"
    ]
    for i, item in enumerate(admin_needs, 1):
        add_numbered_item(doc, i, item)

    add_heading(doc, "3.3 DeepSeek AI服务", 2)
    add_paragraph(doc, "DeepSeek AI服务作为系统外部智能能力提供方，在AI推荐模块中接收系统整理后的用户特征数据，并返回自然语言形式的个性化健身建议。该外部服务不直接访问数据库，而是通过后端服务传入结构化Prompt，从而保证系统边界清晰。")

    # 四
    add_heading(doc, "四、可行性分析", 1)
    add_heading(doc, "4.1 技术可行性", 2)
    add_paragraph(doc, "本系统采用Python语言和Flask框架进行开发，数据库采用SQLite并通过SQLAlchemy进行对象关系映射。前端使用HTML、CSS、Bootstrap和Chart.js完成页面展示和数据可视化。上述技术路线成熟稳定，资料丰富，适合课程设计开发。")
    add_paragraph(doc, "AI推荐模块采用“本地规则模型 + DeepSeek大模型”的混合架构。本地模型用于训练评分、用户画像和基础推荐；DeepSeek用于生成自然语言建议。即使AI接口不可用，系统也可回退到本地模型，保证系统基本功能可用。因此，从技术角度看，本系统具有较高可行性。")

    add_heading(doc, "4.2 时间可行性", 2)
    add_paragraph(doc, "本项目功能范围明确，围绕用户管理、训练饮食记录、身体指标管理、数据分析、训练计划和AI推荐展开。核心功能模块边界清晰，开发任务可按阶段完成，包括需求分析、数据库设计、后端开发、前端页面、AI接口接入和测试优化。")

    add_heading(doc, "4.3 设备与运行可行性", 2)
    add_paragraph(doc, "本系统对硬件要求较低，普通个人电脑即可完成开发与运行。若部署到云端平台，用户可以通过浏览器访问系统。系统使用SQLite作为本地数据库，部署简单，适合课程项目演示。")

    add_heading(doc, "4.4 应用可行性", 2)
    add_paragraph(doc, "大学生群体具有真实的健身与饮食管理需求，同时又普遍缺乏科学训练和饮食指导。本系统将信息管理、数据分析和AI推荐相结合，能够有效满足校园场景中的健康管理需求，具有较强应用可行性。")

    # 五
    add_heading(doc, "五、开发与运行环境", 1)
    add_table(doc, ["类别", "内容"], [
        ["硬件环境", "普通个人电脑，建议8GB内存及以上，5GB以上可用磁盘空间"],
        ["操作系统", "Windows 10/11 或 macOS"],
        ["开发语言", "Python 3.10及以上"],
        ["后端框架", "Flask"],
        ["数据库", "SQLite，Flask-SQLAlchemy"],
        ["前端技术", "HTML、CSS、JavaScript、Bootstrap"],
        ["数据可视化", "Chart.js"],
        ["AI接口", "DeepSeek API"],
        ["开发工具", "VS Code、终端、浏览器"],
        ["运行方式", "本地Flask服务器或云端部署"]
    ])

    # 六
    add_heading(doc, "六、需求分析", 1)
    add_heading(doc, "6.1 学生用户需求", 2)
    add_paragraph(doc, "学生用户希望系统能够提供简洁的注册登录入口，并保存个人基本健康信息。学生希望系统根据自身目标提供训练计划和饮食建议，同时能够记录训练、饮食和身体变化情况。学生还希望系统能够用图表直观展示自己的训练趋势、热量变化和计划完成情况，并在数据不足或执行偏差时给出提醒。")

    add_heading(doc, "6.2 管理员需求", 2)
    add_paragraph(doc, "管理员需要对系统基础数据进行维护，包括学生用户、食物库、系统统计信息等内容。管理员还需要通过后台查看用户数量、训练记录数量、饮食记录数量和计划数量，从整体上掌握系统运行情况。")

    add_heading(doc, "6.3 系统功能需求", 2)
    add_paragraph(doc, "系统主要包括用户登录注册、个人信息管理、训练记录管理、饮食记录管理、身体指标管理、训练计划管理、数据分析、智能推荐、AI大模型推荐和管理员后台管理等功能。")
    add_paragraph(doc, "其中，数据分析功能通过统计训练记录、饮食记录、身体指标和计划完成情况，为学生提供训练质量评分。AI推荐功能通过提取用户行为特征，结合本地模型和DeepSeek大模型，为学生生成个性化建议。")

    add_heading(doc, "6.4 非功能需求", 2)
    add_table(doc, ["需求类别", "具体要求"], [
        ["可用性", "界面清晰，功能入口明确，学生用户可以快速完成记录和查询。"],
        ["安全性", "区分学生用户和管理员角色，限制不同角色访问范围。"],
        ["可靠性", "系统应正确保存训练、饮食和身体指标数据，避免数据丢失。"],
        ["可维护性", "系统采用模块化结构，便于后续扩展推荐模型和功能模块。"],
        ["性能要求", "普通查询、保存和页面展示应在较短时间内完成。"],
        ["扩展性", "后续可扩展云数据库、移动端访问、可穿戴设备数据接入。"]
    ])

    add_heading(doc, "6.5 系统目标", 2)
    for i, item in enumerate([
        "实现学生健康数据的信息化管理。",
        "实现训练、饮食和身体指标的持续记录。",
        "实现基于数据的训练质量评分。",
        "实现训练计划自动生成。",
        "实现DeepSeek大模型智能建议。",
        "实现管理员后台基础数据维护。",
        "体现真实场景下的信息管理、数据分析和AI应用能力。"
    ], 1):
        add_numbered_item(doc, i, item)

    # 七
    add_heading(doc, "七、系统功能需求表", 1)
    add_table(doc, ["功能名称", "详细需求"], [
        ["用户注册与登录", "学生用户可注册账号并登录系统；管理员通过后台入口登录；系统根据角色进入不同页面。"],
        ["个人信息管理", "学生用户可维护年龄、身高、体重、目标等信息，为推荐和分析提供基础。"],
        ["训练记录管理", "用户可新增、查看训练记录，包括训练项目、时长、状态和消耗热量。"],
        ["饮食记录管理", "用户可记录食物名称、摄入量和热量，用于热量摄入分析。"],
        ["身体指标管理", "用户可记录体重、BMI、腰围等指标，并查看历史变化。"],
        ["健身计划管理", "用户可添加、编辑、删除计划，也可一键生成智能训练计划。"],
        ["数据分析功能", "系统计算训练频率、热量收支、计划完成率和训练评分。"],
        ["智能推荐功能", "系统根据目标、BMI、训练频率和热量情况生成本地推荐。"],
        ["DeepSeek AI推荐", "系统调用DeepSeek大模型生成自然语言个性化训练建议。"],
        ["AI降级机制", "当API不可用、余额不足或网络异常时，系统回退至本地推荐模型。"],
        ["管理员后台", "管理员可查看用户、食物库和系统统计数据。"]
    ])

    # 八 数据流图
    add_heading(doc, "八、系统功能建模：分层数据流图", 1)

    add_heading(doc, "8.1 0层数据流图：系统环境图", 2)
    add_paragraph(doc, "0层数据流图用于描述系统边界以及系统与外部实体之间的数据交换关系。本系统外部实体包括学生用户、管理员和DeepSeek AI服务。学生用户向系统提交注册登录信息、个人信息、训练记录、饮食记录和身体指标数据；系统向学生用户返回训练计划、数据分析结果、评分结果和智能推荐建议。管理员向系统提交后台维护请求，并接收后台管理结果。DeepSeek AI服务接收系统整理后的用户特征信息，并返回自然语言推荐建议。")
    add_image_placeholder(
        doc,
        "图8-1 0层数据流图：AI智能健身管理系统环境图",
        "图8-1_0层数据流图.png",
        "建议图中包含：学生用户、管理员、DeepSeek AI服务、AI智能健身管理系统，以及注册信息、训练数据、饮食数据、AI请求、AI建议等数据流。"
    )

    add_heading(doc, "8.2 1层数据流图：系统主要功能分解", 2)
    add_paragraph(doc, "在1层数据流图中，系统被分解为用户管理、数据记录、数据分析、训练计划管理、AI推荐服务和管理员后台六个主要加工过程。用户管理负责注册登录和个人资料维护；数据记录负责训练、饮食和身体指标保存；数据分析负责统计和评分；训练计划管理负责计划增删改查与自动生成；AI推荐服务负责本地模型与DeepSeek模型调用；管理员后台负责基础数据维护。")
    add_image_placeholder(
        doc,
        "图8-2 1层数据流图：系统主要功能分解图",
        "图8-2_1层数据流图.png",
        "建议图中包含：P1用户管理、P2数据记录、P3数据分析、P4训练计划、P5 AI推荐、P6后台管理，以及D1用户库、D2训练记录库、D3饮食记录库、D4身体指标库、D5计划库等数据存储。"
    )

    add_heading(doc, "8.3 2层数据流图：AI推荐子系统", 2)
    add_paragraph(doc, "2层数据流图重点展开AI推荐模块。该模块首先读取用户训练、饮食、身体指标和计划完成情况，然后进行特征提取，形成训练频率、饮食记录完整度、身体指标记录完整度、计划完成率和热量管理等特征。系统通过本地评分模型生成用户画像和基础建议，再判断DeepSeek API是否可用。若API可用，则调用大模型生成自然语言建议；若API不可用，则启用本地推荐结果。最终系统将评分、用户画像、推荐动作和AI文本建议反馈给学生用户。")
    add_image_placeholder(
        doc,
        "图8-3 2层数据流图：AI推荐子系统数据流图",
        "图8-3_2层AI推荐数据流图.png",
        "建议图中包含：用户历史数据、特征提取、本地评分模型、API可用性判断、DeepSeek调用、降级机制、推荐融合、推荐结果。"
    )

    add_heading(doc, "8.4 2层数据流图：训练与饮食记录管理", 2)
    add_paragraph(doc, "训练与饮食记录管理模块进一步分解为训练记录录入、饮食记录录入、历史记录查询、记录修改删除和统计数据输出。该模块是系统数据分析和AI推荐的数据来源，因此需要保证数据录入完整、字段清晰、查询准确。")
    add_image_placeholder(
        doc,
        "图8-4 2层数据流图：训练与饮食记录管理子系统",
        "图8-4_训练饮食记录数据流图.png",
        "建议图中包含：录入训练记录、录入饮食记录、查询历史记录、修改删除记录、训练记录库、饮食记录库。"
    )

    # 九 数据字典
    add_heading(doc, "九、数据字典", 1)

    add_heading(doc, "9.1 数据流词条", 2)
    add_table(doc, ["数据流名", "来源", "去向", "组成", "说明"], [
        ["注册登录信息", "学生用户", "用户管理模块", "username, password", "用于学生用户身份验证。"],
        ["个人基础信息", "学生用户", "用户信息库", "age, height, weight, goal", "用于建立用户健康档案。"],
        ["训练记录信息", "学生用户", "训练记录库", "date, item, duration, calories, status", "用于记录用户训练行为。"],
        ["饮食记录信息", "学生用户", "饮食记录库", "date, food_name, amount, calories", "用于记录用户热量摄入。"],
        ["身体指标信息", "学生用户", "身体指标库", "date, weight, bmi, waist", "用于记录用户身体变化。"],
        ["计划生成请求", "学生用户", "训练计划模块", "goal, user_id", "用于自动生成训练计划。"],
        ["AI推荐请求", "AI推荐模块", "DeepSeek AI服务", "goal, bmi, score, calories, records", "用于生成大模型建议。"],
        ["AI推荐结果", "DeepSeek AI服务", "AI推荐页面", "analysis, advice, weekly_plan", "自然语言智能建议。"],
        ["后台维护信息", "管理员", "后台管理模块", "food, user, statistics", "用于系统数据维护。"]
    ])

    add_heading(doc, "9.2 数据存储词条", 2)
    add_table(doc, ["数据存储名", "编号", "含义", "主要内容"], [
        ["用户信息库", "D1", "保存学生用户信息", "用户编号、账号、密码、年龄、身高、体重、目标"],
        ["管理员信息库", "D2", "保存管理员账号信息", "管理员编号、账号、密码"],
        ["训练记录库", "D3", "保存训练数据", "日期、项目、时长、状态、消耗热量"],
        ["饮食记录库", "D4", "保存饮食数据", "日期、食物、摄入量、热量"],
        ["身体指标库", "D5", "保存身体变化数据", "日期、体重、BMI、腰围"],
        ["食物信息库", "D6", "保存食物热量信息", "食物名称、每100g热量"],
        ["训练计划库", "D7", "保存训练计划数据", "计划名称、动作、组数、次数、重量、完成状态"],
        ["AI推荐结果", "D8", "逻辑数据存储", "特征向量、评分、推荐文本、模型类型"]
    ])

    add_heading(doc, "9.3 加工词条", 2)
    add_table(doc, ["加工名", "编号", "输入", "输出", "加工逻辑"], [
        ["用户登录验证", "P1", "账号、密码", "登录结果", "查询用户表或管理员表，判断账号密码是否匹配。"],
        ["训练记录保存", "P2", "训练项目、时长、热量", "保存结果", "校验训练记录后写入训练记录表。"],
        ["饮食记录保存", "P3", "食物名称、热量", "保存结果", "校验饮食记录后写入饮食记录表。"],
        ["身体指标计算", "P4", "体重、身高", "BMI结果", "根据身高体重计算BMI并保存。"],
        ["训练评分计算", "P5", "训练、饮食、身体指标、计划完成情况", "评分与等级", "按权重计算综合训练质量评分。"],
        ["自动计划生成", "P6", "用户目标", "训练计划", "根据增肌、减脂、塑形目标生成不同训练动作。"],
        ["DeepSeek智能推荐", "P7", "用户特征、评分结果、本地建议", "AI建议", "构造Prompt并调用DeepSeek大模型。"],
        ["AI降级处理", "P8", "API状态", "本地推荐结果", "当API不可用时返回本地模型建议。"]
    ])

    add_heading(doc, "9.4 详细数据项字典", 2)
    data_items = {
        "User 用户信息表": [
            ["id", "用户编号", "Integer", "主键，自增"],
            ["username", "用户名", "String(50)", "唯一，不能为空"],
            ["password", "密码", "String(100)", "不能为空"],
            ["age", "年龄", "Integer", "可为空"],
            ["height", "身高", "Float", "单位cm"],
            ["weight", "体重", "Float", "单位kg"],
            ["goal", "健身目标", "String(50)", "增肌、减脂、塑形等"]
        ],
        "Admin 管理员表": [
            ["id", "管理员编号", "Integer", "主键，自增"],
            ["username", "管理员账号", "String(50)", "唯一，不能为空"],
            ["password", "管理员密码", "String(100)", "不能为空"]
        ],
        "WorkoutRecord 训练记录表": [
            ["id", "训练记录编号", "Integer", "主键，自增"],
            ["date", "训练日期", "Date", "默认当天"],
            ["item", "训练项目", "String(100)", "不能为空"],
            ["duration", "训练时长", "Integer", "单位分钟"],
            ["status", "完成状态", "String(50)", "默认已完成"],
            ["calories", "消耗热量", "Integer", "默认0"],
            ["user_id", "用户编号", "Integer", "外键"]
        ],
        "DietRecord 饮食记录表": [
            ["id", "饮食记录编号", "Integer", "主键，自增"],
            ["date", "记录日期", "Date", "默认当天"],
            ["food_name", "食物名称", "String(100)", "不能为空"],
            ["amount", "摄入量", "String(50)", "可为空"],
            ["calories", "摄入热量", "Integer", "默认0"],
            ["user_id", "用户编号", "Integer", "外键"]
        ],
        "BodyRecord 身体指标表": [
            ["id", "身体记录编号", "Integer", "主键，自增"],
            ["date", "记录日期", "Date", "默认当天"],
            ["weight", "体重", "Float", "不能为空"],
            ["bmi", "身体质量指数", "Float", "由系统计算"],
            ["waist", "腰围", "Float", "可为空"],
            ["user_id", "用户编号", "Integer", "外键"]
        ],
        "Food 食物库表": [
            ["id", "食物编号", "Integer", "主键，自增"],
            ["name", "食物名称", "String(100)", "不能为空"],
            ["calories_per_100g", "每100g热量", "Integer", "不能为空"]
        ],
        "WorkoutPlan 健身计划表": [
            ["id", "计划编号", "Integer", "主键，自增"],
            ["plan_name", "计划名称", "String(100)", "不能为空"],
            ["action", "训练动作", "String(100)", "不能为空"],
            ["sets", "组数", "Integer", "可为空"],
            ["reps", "次数", "Integer", "可为空"],
            ["weight", "重量", "Float", "可为空"],
            ["is_completed", "是否完成", "Boolean", "默认False"],
            ["user_id", "用户编号", "Integer", "外键"]
        ]
    }

    for table_name, rows in data_items.items():
        add_heading(doc, table_name, 2)
        add_table(doc, ["字段名", "含义", "类型", "约束/说明"], rows)

    # 十 E-R
    add_heading(doc, "十、系统数据建模：E-R模型", 1)
    add_heading(doc, "10.1 核心实体", 2)
    for i, item in enumerate([
        "User（用户）：表示学生用户，是系统的核心实体。",
        "Admin（管理员）：表示后台管理人员。",
        "WorkoutRecord（训练记录）：表示用户每次训练行为。",
        "DietRecord（饮食记录）：表示用户每日饮食摄入。",
        "BodyRecord（身体指标）：表示用户体重、BMI等身体变化。",
        "Food（食物）：表示系统食物库中的基础热量数据。",
        "WorkoutPlan（健身计划）：表示用户训练计划与执行状态。",
        "AIRecommendation（AI推荐结果）：表示系统根据用户特征生成的推荐结果，可作为逻辑实体。"
    ], 1):
        add_numbered_item(doc, i, item)

    add_heading(doc, "10.2 实体关系说明", 2)
    relation_texts = [
        "一个用户可以对应多条训练记录，User 与 WorkoutRecord 为 1:N 关系。",
        "一个用户可以对应多条饮食记录，User 与 DietRecord 为 1:N 关系。",
        "一个用户可以对应多条身体指标记录，User 与 BodyRecord 为 1:N 关系。",
        "一个用户可以拥有多条健身计划，User 与 WorkoutPlan 为 1:N 关系。",
        "一种食物可以被多条饮食记录引用，Food 与 DietRecord 可设计为 1:N 关系。",
        "一个用户可以产生多次AI推荐结果，User 与 AIRecommendation 为 1:N 关系。",
        "管理员可以维护食物库和系统基础数据，Admin 与 Food、WorkoutPlan 存在维护关系。"
    ]
    for i, item in enumerate(relation_texts, 1):
        add_numbered_item(doc, i, item)

    add_image_placeholder(
        doc,
        "图10-1 系统E-R模型图",
        "图10-1_ER模型图.png",
        "建议图中包含：User、Admin、WorkoutRecord、DietRecord、BodyRecord、Food、WorkoutPlan、AIRecommendation，并标注1:N关系。"
    )

    # 十一 关系模式
    add_heading(doc, "十一、关系模式", 1)
    relations = [
        "User(id, username, password, age, height, weight, goal)",
        "Admin(id, username, password)",
        "WorkoutRecord(id, date, item, duration, status, calories, user_id)",
        "DietRecord(id, date, food_name, amount, calories, user_id)",
        "BodyRecord(id, date, weight, bmi, waist, user_id)",
        "Food(id, name, calories_per_100g)",
        "WorkoutPlan(id, plan_name, action, sets, reps, weight, is_completed, user_id)",
        "AIRecommendation(id, user_id, feature_vector, local_score, model_result, deepseek_advice, created_at)"
    ]
    for r in relations:
        add_paragraph(doc, r, first_line=False)

    add_paragraph(doc, "其中，WorkoutRecord.user_id、DietRecord.user_id、BodyRecord.user_id、WorkoutPlan.user_id均为外键，关联User.id。AIRecommendation为逻辑关系模式，用于描述AI推荐结果的存储设计，可在后续系统扩展时落地为实际数据库表。")

    # 十二
    add_heading(doc, "十二、非功能需求", 1)
    add_table(doc, ["需求类别", "具体描述"], [
        ["安全性", "系统应区分学生用户和管理员角色，避免越权访问。"],
        ["易用性", "页面布局应简洁，按钮和导航清晰，便于用户快速操作。"],
        ["可靠性", "系统应稳定保存训练、饮食和身体指标数据。"],
        ["可维护性", "推荐逻辑、数据分析逻辑和AI调用逻辑应保持模块化，便于维护。"],
        ["可扩展性", "系统可扩展云端部署、移动端适配、外部设备数据接入和更多AI模型。"],
        ["可解释性", "AI推荐结果应结合本地评分模型，使用户理解推荐依据。"]
    ])

    # 十三
    add_heading(doc, "十三、创新点总结", 1)
    innovations = [
        "本系统不是传统CRUD系统，而是将信息管理、数据分析和AI推荐结合的智能系统。",
        "系统构建训练质量评分模型，将用户行为数据转化为可量化分数。",
        "系统提取训练频率、饮食记录、身体指标、计划完成率和热量平衡等特征，体现特征工程思想。",
        "系统接入DeepSeek大模型，实现自然语言个性化健身建议。",
        "系统设计AI降级机制，当API不可用时自动回退到本地推荐模型。",
        "系统具备模块化服务层扩展潜力，可进一步演进为微服务架构。"
    ]
    for i, item in enumerate(innovations, 1):
        add_numbered_item(doc, i, item)

    # 十四
    add_heading(doc, "十四、项目组成员分工", 1)
    add_table(doc, ["成员", "主要负责模块", "具体工作"], [
        ["请填写", "后端与数据库", "Flask路由设计、SQLAlchemy模型设计、数据库初始化、用户登录与记录管理。"],
        ["请填写", "前端与可视化", "页面设计、Bootstrap样式、Chart.js图表、用户交互优化。"],
        ["请填写", "AI与文档", "训练评分模型、DeepSeek API接入、需求文档、数据流图和E-R图设计。"]
    ])

    doc.save(DOCX_PATH)
    print(f"需求分析文档已生成：{DOCX_PATH.resolve()}")
    print(f"图片文件夹路径：{IMAGE_DIR.resolve()}")
    print("如果要插入手绘图片，请按文档中的图片文件名保存到 images 文件夹后重新运行。")


if __name__ == "__main__":
    build_doc()