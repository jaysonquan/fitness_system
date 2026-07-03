# -*- coding: utf-8 -*-
"""
AI智能健身管理系统
概要设计与详细设计文档自动生成脚本

运行前安装：
pip install python-docx

生成结果：
design_output/AI智能健身管理系统_概要设计与详细设计文档.docx

图片放置规则：
将图片放入 images_design 文件夹，并按文档中提示的文件名命名。
再次运行脚本后，图片会自动插入文档。
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_DIR = Path("design_output")
IMAGE_DIR = Path("images_design")
OUTPUT_DIR.mkdir(exist_ok=True)
IMAGE_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUTPUT_DIR / "AI智能健身管理系统_概要设计与详细设计文档.docx"


# =========================
# 基础格式函数
# =========================

def set_run_font(run, font_name="宋体", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph(paragraph, first_line=True, center=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(6)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def add_cover_line(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, "黑体" if bold else "宋体", size, bold)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, "黑体", 16 if level == 1 else 14, True)


def add_para(doc, text, first_line=True):
    p = doc.add_paragraph()
    set_paragraph(p, first_line=first_line)
    r = p.add_run(text)
    set_run_font(r, "宋体", 12, False)


def add_number(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(f"{num}. {text}")
    set_run_font(r, "宋体", 12, False)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_run_font(r, "宋体", 10, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False)

    doc.add_paragraph()
    return table


def add_image_placeholder(doc, title, filename, description):
    add_para(doc, f"【图片放置位置：{title}】", first_line=False)

    image_path = IMAGE_DIR / filename
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.8))
    else:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.text = ""

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p.add_run(title + "\n")
        set_run_font(r1, "黑体", 12, True)

        r2 = p.add_run(description + "\n")
        set_run_font(r2, "宋体", 11, False)

        r3 = p.add_run(f"请将图片命名为：{filename}，放入 images_design 文件夹后重新运行脚本。")
        set_run_font(r3, "宋体", 10, False)

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run(title)
    set_run_font(r, "黑体", 11, True)


def add_mermaid_block(doc, title, code):
    add_para(doc, f"【{title} Mermaid参考代码】", first_line=False)
    p = doc.add_paragraph()
    r = p.add_run(code)
    set_run_font(r, "宋体", 9, False)


# =========================
# 文档主体
# =========================

def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    # 封面
    add_cover_line(doc, "《数据应用与实践》课程", 16, True)
    doc.add_paragraph()
    add_cover_line(doc, "概要设计与详细设计文档", 22, True)
    doc.add_paragraph()
    add_cover_line(doc, "AI智能健身管理系统", 24, True)
    doc.add_paragraph()
    add_cover_line(doc, "——面向大学生的健身与饮食辅助管理信息系统", 16, True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_cover_line(doc, "课程名称：数据应用与实践", 14)
    add_cover_line(doc, "项目名称：AI智能健身管理系统", 14)
    add_cover_line(doc, "班级：请自行填写", 14)
    add_cover_line(doc, "姓名：请自行填写", 14)
    add_cover_line(doc, "学号：请自行填写", 14)
    add_cover_line(doc, "指导教师：请自行填写", 14)
    add_cover_line(doc, "完成日期：2026年5月", 14)
    doc.add_page_break()

    # 目录
    add_heading(doc, "目录", 1)
    toc = [
        "一、引言",
        "二、前期需求分析文档分析与数据流图改进",
        "三、系统体系架构设计",
        "四、由数据流图转换得到系统功能结构图",
        "五、接口设计",
        "六、数据设计",
        "七、过程设计",
        "八、运行设计",
        "九、出错处理与系统维护设计",
        "十、阶段小结",
        "附录：主要图表清单"
    ]
    for item in toc:
        add_para(doc, item, first_line=False)
    doc.add_page_break()

    # 一、引言
    add_heading(doc, "一、引言", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_para(doc, "本概要设计与详细设计文档是在前期需求分析文档的基础上，对AI智能健身管理系统进行进一步的软件设计说明。本文档的主要目的在于明确系统的体系结构、功能模块划分、模块接口、数据库设计和核心算法流程，为后续系统编码、测试、部署和答辩展示提供依据。")
    add_para(doc, "与需求分析阶段主要关注“系统需要做什么”不同，概要设计与详细设计阶段更加关注“系统如何实现”。因此，本文档将对前期数据流图进行分析与改进，并在此基础上通过变换分析产生系统功能结构图，同时完成接口设计、数据设计和过程设计。")

    add_heading(doc, "1.2 系统范围", 2)
    add_para(doc, "本系统面向大学生健身与饮食管理场景，主要支持学生用户完成个人信息维护、训练记录、饮食记录、身体指标记录、训练计划管理、数据分析、训练质量评分和AI智能推荐等操作；管理员则主要负责用户管理、食物库维护和系统统计查看。系统引入DeepSeek大模型接口，在本地规则模型基础上生成个性化训练建议。")

    add_heading(doc, "1.3 设计原则", 2)
    principles = [
        "高内聚、低耦合：将用户管理、记录管理、数据分析、训练计划和AI推荐划分为相对独立的模块。",
        "分层设计：采用表示层、业务逻辑层、数据访问层和外部AI服务层的分层结构。",
        "数据驱动：以用户训练、饮食和身体指标数据为基础，实现统计分析和智能推荐。",
        "可靠性优先：AI接口不可用时，系统自动回退到本地推荐模型，保证核心功能不受影响。",
        "可扩展性：系统后续可扩展云数据库、移动端、更多AI模型和可穿戴设备数据接入。"
    ]
    for i, item in enumerate(principles, 1):
        add_number(doc, i, item)

    # 二、需求分析与DFD改进
    add_heading(doc, "二、前期需求分析文档分析与数据流图改进", 1)
    add_heading(doc, "2.1 前期需求分析文档存在的问题", 2)
    add_para(doc, "前期需求分析文档已经完成了系统背景、用户说明、功能需求、数据字典、E-R模型和关系模式等内容，基本明确了系统功能边界。但从结构化设计角度看，前期数据流图仍存在进一步完善空间。")
    issues = [
        "部分数据流命名较为宽泛，如“数据”“信息”等，缺少明确业务含义。",
        "AI推荐模块在前期图中已经出现，但内部处理流程不够细化。",
        "数据存储与处理过程之间的对应关系需要进一步明确。",
        "管理员后台维护与普通用户业务流程应进一步区分。",
        "训练、饮食、身体指标记录之间的数据汇总关系需要加强表达。"
    ]
    for i, item in enumerate(issues, 1):
        add_number(doc, i, item)

    add_heading(doc, "2.2 数据流图改进原则", 2)
    rules = [
        "将外部实体、处理过程和数据存储分开表示，避免图中元素混杂。",
        "数据流名称尽量使用业务名词，如训练记录数据、饮食记录数据、用户特征数据、AI建议等。",
        "突出AI推荐子系统的数据处理链路，包括特征提取、本地模型、API判断、DeepSeek调用和降级机制。",
        "将后台管理流程与学生用户使用流程分离，提高系统结构清晰度。",
        "将数据记录模块细化为训练记录、饮食记录、历史查询和修改删除处理。"
    ]
    for i, item in enumerate(rules, 1):
        add_number(doc, i, item)

    add_heading(doc, "2.3 改进后的0层数据流图", 2)
    add_para(doc, "0层数据流图用于描述系统边界及系统与外部实体之间的数据交换关系。改进后的0层数据流图应包括学生用户、管理员、DeepSeek AI服务和AI智能健身管理系统四类主要对象。")
    add_image_placeholder(
        doc,
        "图2-1 改进后的0层数据流图：AI智能健身管理系统环境图",
        "图2-1_改进0层数据流图.png",
        "建议图中包含：学生用户、管理员、DeepSeek AI服务、AI智能健身管理系统，以及注册信息、训练数据、饮食数据、AI请求、AI建议等数据流。"
    )

    add_heading(doc, "2.4 改进后的1层数据流图", 2)
    add_para(doc, "1层数据流图在0层图基础上，对系统内部主要处理过程进行分解。改进后的1层数据流图应包括P1用户管理、P2数据记录、P3数据分析、P4训练计划、P5 AI推荐、P6后台管理等处理过程，并对应D1用户库、D2训练记录库、D3饮食记录库、D4身体指标库、D5计划库等数据存储。")
    add_image_placeholder(
        doc,
        "图2-2 改进后的1层数据流图：系统主要功能分解图",
        "图2-2_改进1层数据流图.png",
        "建议图中包含：P1用户管理、P2数据记录、P3数据分析、P4训练计划、P5 AI推荐、P6后台管理，以及D1用户库、D2训练记录库、D3饮食记录库、D4身体指标库、D5计划库。"
    )

    add_heading(doc, "2.5 改进后的2层数据流图：AI推荐子系统", 2)
    add_para(doc, "AI推荐子系统是本系统的创新模块。改进后的2层数据流图应重点描述从用户历史数据到最终推荐结果的处理过程，包括用户历史数据汇总、特征提取、本地评分模型、API可用性判断、DeepSeek调用、降级机制、推荐融合和推荐结果输出。")
    add_image_placeholder(
        doc,
        "图2-3 改进后的2层数据流图：AI推荐子系统数据流图",
        "图2-3_改进AI推荐数据流图.png",
        "建议图中包含：用户历史数据、特征提取、本地评分模型、API可用性判断、DeepSeek调用、降级机制、推荐融合、推荐结果。"
    )

    add_heading(doc, "2.6 改进后的2层数据流图：训练与饮食记录管理", 2)
    add_para(doc, "训练与饮食记录管理模块是系统数据来源。改进后的图应体现用户录入训练记录、录入饮食记录、查询历史记录和修改删除记录的流程，并分别连接训练记录库和饮食记录库。")
    add_image_placeholder(
        doc,
        "图2-4 改进后的2层数据流图：训练与饮食记录管理子系统",
        "图2-4_改进训练饮食记录数据流图.png",
        "建议图中包含：录入训练记录、录入饮食记录、查询历史记录、修改删除记录、训练记录库、饮食记录库。"
    )

    # 三、架构设计
    add_heading(doc, "三、系统体系架构设计", 1)
    add_heading(doc, "3.1 架构设计思想", 2)
    add_para(doc, "本系统采用B/S架构和分层设计思想。用户通过浏览器访问系统页面，前端页面向Flask后端提交请求，后端完成业务处理和数据库操作，并在AI推荐场景下调用DeepSeek外部AI服务。系统整体可划分为表示层、业务逻辑层、数据访问层、数据库层和外部AI服务层。")

    add_heading(doc, "3.2 系统体系架构示意图", 2)
    add_image_placeholder(
        doc,
        "图3-1 系统体系架构示意图",
        "图3-1_系统体系架构示意图.png",
        "建议图中包含：浏览器/用户端、Flask Web应用、业务服务层、数据库SQLite、DeepSeek API。可采用四层结构：表示层、业务逻辑层、数据访问层、外部AI服务层。"
    )

    add_heading(doc, "3.3 分层结构说明", 2)
    add_table(doc, ["层次", "组成", "主要职责"], [
        ["表示层", "HTML、CSS、Bootstrap、Chart.js", "负责页面展示、表单输入、图表展示和用户交互。"],
        ["控制层", "Flask路由", "接收请求、进行登录检查、调用业务逻辑并返回页面。"],
        ["业务逻辑层", "用户管理、记录管理、评分模型、计划生成、AI推荐", "实现系统核心业务规则和推荐算法。"],
        ["数据访问层", "SQLAlchemy模型", "完成数据库表映射、查询、插入、更新和删除操作。"],
        ["数据存储层", "SQLite数据库", "保存用户、训练记录、饮食记录、身体指标、训练计划和食物库数据。"],
        ["外部AI服务层", "DeepSeek API", "根据用户特征生成自然语言个性化建议。"]
    ])

    add_heading(doc, "3.4 模块划分", 2)
    add_table(doc, ["模块名称", "主要功能", "对应页面或路由"], [
        ["用户管理模块", "注册、登录、退出、个人资料维护", "/register, /login, /logout, /profile"],
        ["训练记录模块", "新增训练记录、查看训练历史", "/workouts, /workouts/add"],
        ["饮食记录模块", "新增饮食记录、查看饮食历史", "/diets, /diets/add"],
        ["身体指标模块", "新增身体指标、查看体重趋势", "/bodies, /bodies/add"],
        ["训练计划模块", "新增计划、编辑计划、删除计划、完成计划、自动生成计划", "/workout_plans, /plans/add, /plans/generate"],
        ["数据分析模块", "统计训练次数、热量摄入、热量消耗、体重变化和完成率", "/dashboard, /score"],
        ["AI推荐模块", "本地评分模型、DeepSeek调用、降级推荐", "/ai_recommend"],
        ["管理员模块", "管理员登录、用户查看、食物库维护、系统统计", "/admin/login, /admin/dashboard, /admin/users, /admin/foods"]
    ])

    # 四、功能结构转换
    add_heading(doc, "四、由数据流图转换得到系统功能结构图", 1)
    add_heading(doc, "4.1 转换方法说明", 2)
    add_para(doc, "结构化设计中，可根据数据流图进行变换分析，将数据流图中的输入流、中心处理和输出流转化为软件结构图。对本系统而言，学生用户输入训练、饮食、身体指标等数据，系统经过记录、统计、评分、推荐等处理，最终输出数据分析结果、训练计划和AI建议。因此系统功能结构图可由数据流图转换得到。")

    add_heading(doc, "4.2 第一步：最初转换结果图", 2)
    add_para(doc, "第一步转换时，直接将数据流图中的主要处理过程映射为功能模块，因此得到的结构较为接近数据流图，模块划分较粗，主要包括用户管理、数据记录、数据分析、训练计划、AI推荐和后台管理。")
    add_image_placeholder(
        doc,
        "图4-1 最初的转换结果图",
        "图4-1_最初转换结果图.png",
        "建议图中包含：AI智能健身管理系统作为根节点，下分用户管理、数据记录、数据分析、训练计划、AI推荐、后台管理六个一级模块。"
    )

    add_heading(doc, "4.3 第二步：最终改进优化后的功能结构图", 2)
    add_para(doc, "第二步优化时，对第一步得到的功能结构图进一步分解和调整，使模块职责更加清晰。将数据记录模块拆分为训练记录、饮食记录和身体指标记录；将数据分析模块拆分为热量分析、体重趋势、训练评分和计划完成率；将AI推荐模块拆分为特征提取、本地模型、DeepSeek调用和降级机制。")
    add_image_placeholder(
        doc,
        "图4-2 最终改进优化后的功能结构图",
        "图4-2_最终功能结构图.png",
        "建议图中包含：用户管理、记录管理、数据分析、计划管理、AI推荐、管理员后台六大模块，并继续细化二级功能。"
    )

    add_heading(doc, "4.4 功能结构优化说明", 2)
    add_table(doc, ["优化前模块", "存在问题", "优化后设计"], [
        ["数据记录", "训练、饮食、身体指标混在一起", "拆分为训练记录管理、饮食记录管理、身体指标管理。"],
        ["数据分析", "分析内容不够明确", "拆分为热量分析、体重趋势、计划完成率和训练评分。"],
        ["AI推荐", "只表现为单一模块", "细化为特征提取、本地评分模型、DeepSeek调用和降级机制。"],
        ["后台管理", "功能边界不清", "明确为用户管理、食物库维护和系统统计。"]
    ])

    # 五、接口设计
    add_heading(doc, "五、接口设计", 1)
    add_heading(doc, "5.1 人机界面设计", 2)
    add_para(doc, "人机界面设计主要描述用户通过浏览器与系统进行交互的主要页面。系统界面应保持简洁、清晰、统一，主要采用导航栏方式组织功能入口，使用户能够快速进入首页、个人信息、训练、饮食、身体指标、计划、评分和AI推荐等功能页面。")

    ui_pages = [
        ["登录界面", "输入用户名、密码，提交后进行身份验证；学生用户进入系统首页，管理员进入后台首页。", "图5-1_登录界面设计.png"],
        ["系统首页/仪表盘", "展示训练次数、饮食记录、热量摄入、热量消耗、体重变化和计划完成率。", "图5-2_系统首页界面设计.png"],
        ["训练记录界面", "支持新增训练记录，显示训练历史，包括训练项目、时长、完成状态和消耗热量。", "图5-3_训练记录界面设计.png"],
        ["饮食记录界面", "支持新增饮食记录，显示食物名称、摄入量和摄入热量。", "图5-4_饮食记录界面设计.png"],
        ["身体指标界面", "支持录入体重、腰围等指标，并展示体重趋势图。", "图5-5_身体指标界面设计.png"],
        ["训练计划界面", "支持创建、编辑、删除、完成和自动生成训练计划。", "图5-6_训练计划界面设计.png"],
        ["AI推荐界面", "展示本地评分、用户画像、推荐动作和DeepSeek生成的自然语言建议。", "图5-7_AI推荐界面设计.png"],
        ["管理员后台界面", "展示用户数量、训练记录数量、食物库管理和系统统计信息。", "图5-8_管理员后台界面设计.png"]
    ]

    for title, desc, filename in ui_pages:
        add_heading(doc, title, 2)
        add_para(doc, desc)
        add_image_placeholder(
            doc,
            f"{title}示意图",
            filename,
            f"请放置{title}截图或界面原型图。"
        )

    add_heading(doc, "5.2 主要功能模块接口设计", 2)
    add_table(doc, ["模块", "接口/路由", "输入", "输出", "说明"], [
        ["用户登录", "POST /login", "username, password", "登录结果/跳转页面", "验证学生账号密码。"],
        ["用户注册", "POST /register", "username, password", "注册结果", "创建新学生用户。"],
        ["个人信息", "POST /profile", "age, height, weight, goal", "保存结果", "维护用户基本信息。"],
        ["训练记录", "POST /workouts/add", "date, item, duration, calories, status", "训练记录列表", "新增训练记录。"],
        ["饮食记录", "POST /diets/add", "date, food_name, amount, calories", "饮食记录列表", "新增饮食记录。"],
        ["身体指标", "POST /bodies/add", "date, weight, waist", "身体指标列表", "计算BMI并保存。"],
        ["训练计划", "GET /plans/generate", "user.goal", "训练计划列表", "根据用户目标自动生成训练计划。"],
        ["评分系统", "GET /score", "用户历史数据", "评分、等级、评价", "计算训练质量得分。"],
        ["AI推荐", "GET /ai_recommend", "用户画像、评分、历史数据", "推荐建议", "调用本地模型和DeepSeek模型。"],
        ["管理员食物库", "POST /admin/foods", "food_name, calories", "食物库列表", "新增食物热量数据。"]
    ])

    add_heading(doc, "5.3 外部接口设计", 2)
    add_table(doc, ["外部接口", "调用方式", "输入", "输出", "异常处理"], [
        ["DeepSeek API", "OpenAI兼容接口", "用户目标、BMI、训练频率、饮食记录、评分结果", "自然语言训练建议", "失败时返回本地推荐结果。"],
        ["浏览器", "HTTP请求", "表单数据、页面请求", "HTML页面、图表数据", "输入错误时返回提示信息。"],
        ["SQLite数据库", "SQLAlchemy ORM", "查询、插入、更新、删除请求", "数据库记录", "数据库异常时提示操作失败。"]
    ])

    add_heading(doc, "5.4 内部接口设计", 2)
    add_table(doc, ["调用方", "被调用模块", "传入数据", "返回数据", "说明"], [
        ["路由层", "用户服务", "账号、密码、个人资料", "用户对象/验证结果", "完成用户注册登录与资料维护。"],
        ["路由层", "记录服务", "训练/饮食/身体指标表单", "保存结果", "完成用户数据记录。"],
        ["数据分析模块", "数据库模型", "user_id", "统计结果", "读取用户历史数据用于分析。"],
        ["AI推荐模块", "本地评分模型", "用户特征", "评分和用户画像", "生成基础推荐。"],
        ["AI推荐模块", "DeepSeek服务", "Prompt文本", "AI建议", "生成自然语言建议。"]
    ])

    # 六、数据设计
    add_heading(doc, "六、数据设计", 1)
    add_heading(doc, "6.1 数据库设计说明", 2)
    add_para(doc, "本系统数据库以用户表为核心，训练记录、饮食记录、身体指标、训练计划和AI推荐均通过user_id与用户表建立关联。管理员表用于后台登录；食物库表用于保存基础热量数据。数据库设计遵循实体完整性和参照完整性原则。")

    database_tables = {
        "6.2 用户信息表 User": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["username", "String", "50", "用户名，唯一，不能为空"],
            ["password", "String", "100", "密码，不能为空"],
            ["age", "Integer", "3", "年龄，可为空"],
            ["height", "Float", "-", "身高，单位cm"],
            ["weight", "Float", "-", "当前体重，单位kg"],
            ["goal", "String", "50", "健身目标，如增肌、减脂、塑形"]
        ],
        "6.3 管理员表 Admin": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["username", "String", "50", "管理员账号，不能为空"],
            ["password", "String", "100", "管理员密码，不能为空"]
        ],
        "6.4 训练记录表 WorkoutRecord": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["user_id", "Integer", "11", "外键，关联User.id"],
            ["date", "Date", "-", "训练日期"],
            ["item", "String", "100", "训练项目"],
            ["duration", "Integer", "4", "训练时长，单位分钟"],
            ["status", "String", "50", "完成状态"],
            ["calories", "Integer", "6", "消耗热量"]
        ],
        "6.5 饮食记录表 DietRecord": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["user_id", "Integer", "11", "外键，关联User.id"],
            ["date", "Date", "-", "记录日期"],
            ["food_name", "String", "100", "食物名称"],
            ["amount", "String", "50", "摄入量"],
            ["calories", "Integer", "6", "摄入热量"]
        ],
        "6.6 身体指标表 BodyRecord": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["user_id", "Integer", "11", "外键，关联User.id"],
            ["date", "Date", "-", "记录日期"],
            ["weight", "Float", "-", "体重"],
            ["bmi", "Float", "-", "身体质量指数，由系统计算"],
            ["waist", "Float", "-", "腰围"]
        ],
        "6.7 食物库表 Food": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["name", "String", "100", "食物名称"],
            ["calories_per_100g", "Integer", "6", "每100g热量"]
        ],
        "6.8 训练计划表 WorkoutPlan": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["user_id", "Integer", "11", "外键，关联User.id"],
            ["plan_name", "String", "100", "计划名称"],
            ["action", "String", "100", "训练动作"],
            ["sets", "Integer", "3", "组数"],
            ["reps", "Integer", "4", "次数"],
            ["weight", "Float", "-", "训练重量"],
            ["is_completed", "Boolean", "-", "是否完成"]
        ],
        "6.9 AI推荐逻辑表 AIRecommendation": [
            ["字段名", "类型", "长度", "约束/说明"],
            ["id", "Integer", "11", "主键，自增"],
            ["user_id", "Integer", "11", "外键，关联User.id"],
            ["feature_vector", "Text", "-", "用户特征向量"],
            ["local_score", "Float", "-", "本地评分模型得分"],
            ["model_result", "String", "100", "用户画像类型"],
            ["deepseek_advice", "Text", "-", "DeepSeek生成建议"],
            ["created_at", "DateTime", "-", "生成时间"]
        ]
    }

    for table_title, rows in database_tables.items():
        add_heading(doc, table_title, 2)
        add_table(doc, rows[0], rows[1:])

    add_heading(doc, "6.10 数据表关系说明", 2)
    relation_desc = [
        "User与WorkoutRecord为1:N关系，一个用户可以拥有多条训练记录。",
        "User与DietRecord为1:N关系，一个用户可以拥有多条饮食记录。",
        "User与BodyRecord为1:N关系，一个用户可以拥有多条身体指标记录。",
        "User与WorkoutPlan为1:N关系，一个用户可以拥有多条训练计划。",
        "User与AIRecommendation为1:N关系，一个用户可以多次生成AI推荐结果。",
        "DietRecord可与Food建立逻辑关联，用于根据食物库热量数据辅助计算摄入热量。"
    ]
    for i, item in enumerate(relation_desc, 1):
        add_number(doc, i, item)

    add_image_placeholder(
        doc,
        "图6-1 系统数据库关系图",
        "图6-1_系统数据库关系图.png",
        "建议图中包含：User、Admin、WorkoutRecord、DietRecord、BodyRecord、Food、WorkoutPlan、AIRecommendation，并标注1:N关系。"
    )

    # 七、过程设计
    add_heading(doc, "七、过程设计", 1)
    add_heading(doc, "7.1 用户登录流程", 2)
    add_para(doc, "用户登录流程用于验证学生用户或管理员身份。系统根据不同登录入口查询不同数据表，若账号密码正确则写入session并跳转对应页面，否则返回错误提示。")
    add_image_placeholder(
        doc,
        "图7-1 用户登录流程图",
        "图7-1_用户登录流程图.png",
        "流程建议：开始 → 输入账号密码 → 查询用户表/管理员表 → 判断是否匹配 → 登录成功/失败提示 → 结束。"
    )

    add_heading(doc, "7.2 添加训练记录流程", 2)
    add_para(doc, "添加训练记录流程用于保存学生用户的训练行为。用户填写训练日期、项目、时长、状态和消耗热量，系统校验后写入训练记录表。")
    add_image_placeholder(
        doc,
        "图7-2 添加训练记录流程图",
        "图7-2_添加训练记录流程图.png",
        "流程建议：开始 → 登录校验 → 输入训练信息 → 数据合法性检查 → 写入训练记录表 → 返回训练记录列表 → 结束。"
    )

    add_heading(doc, "7.3 添加饮食记录流程", 2)
    add_para(doc, "添加饮食记录流程用于保存用户每日饮食摄入情况。用户填写食物名称、摄入量和热量后，系统保存到饮食记录表，为热量分析和AI推荐提供数据基础。")
    add_image_placeholder(
        doc,
        "图7-3 添加饮食记录流程图",
        "图7-3_添加饮食记录流程图.png",
        "流程建议：开始 → 登录校验 → 输入饮食信息 → 校验热量数据 → 写入饮食记录表 → 返回饮食记录列表 → 结束。"
    )

    add_heading(doc, "7.4 身体指标记录与BMI计算流程", 2)
    add_para(doc, "身体指标记录流程用于保存用户体重、腰围等身体变化数据。当用户身高信息存在时，系统根据体重和身高自动计算BMI，并将结果写入身体指标表。")
    add_image_placeholder(
        doc,
        "图7-4 身体指标记录与BMI计算流程图",
        "图7-4_BMI计算流程图.png",
        "流程建议：开始 → 输入体重和腰围 → 读取用户身高 → 判断身高是否存在 → 计算BMI → 保存身体指标 → 结束。"
    )

    add_heading(doc, "7.5 自动生成训练计划流程", 2)
    add_para(doc, "自动生成训练计划流程根据用户目标匹配计划模板。若用户目标为增肌，则生成力量训练计划；若目标为减脂，则生成有氧和HIIT计划；若目标为塑形，则生成力量与有氧结合计划。")
    add_image_placeholder(
        doc,
        "图7-5 自动生成训练计划流程图",
        "图7-5_自动生成训练计划流程图.png",
        "流程建议：开始 → 读取用户目标 → 判断目标类型 → 匹配计划模板 → 写入训练计划表 → 返回计划列表 → 结束。"
    )

    add_heading(doc, "7.6 训练质量评分算法流程", 2)
    add_para(doc, "训练质量评分算法综合考虑训练频率、饮食记录、身体指标记录、计划完成度和热量管理五个维度。每个维度设置不同权重，最终得到总分和评价等级。")
    add_table(doc, ["评分维度", "权重", "计算依据"], [
        ["训练频率", "30分", "训练记录数量，最多计30分。"],
        ["饮食记录", "20分", "饮食记录数量，最多计20分。"],
        ["身体指标", "15分", "身体指标记录数量，最多计15分。"],
        ["计划完成度", "25分", "已完成计划数/总计划数。"],
        ["热量管理", "10分", "根据用户目标判断摄入与消耗是否匹配。"]
    ])
    add_image_placeholder(
        doc,
        "图7-6 训练质量评分算法流程图",
        "图7-6_训练质量评分流程图.png",
        "流程建议：开始 → 读取历史数据 → 计算各维度分数 → 加权求和 → 判断等级 → 输出评价 → 结束。"
    )

    add_heading(doc, "7.7 AI智能推荐算法流程", 2)
    add_para(doc, "AI智能推荐算法是本系统核心创新流程。系统首先读取用户历史数据并提取特征，随后通过本地规则模型生成评分和基础建议；接着判断DeepSeek API是否可用。若可用，则调用DeepSeek生成自然语言建议；若不可用，则使用本地模型结果作为降级输出。")
    add_image_placeholder(
        doc,
        "图7-7 AI智能推荐算法流程图",
        "图7-7_AI智能推荐流程图.png",
        "流程建议：开始 → 读取历史数据 → 特征提取 → 本地评分模型 → 判断API可用性 → DeepSeek调用/本地降级 → 推荐融合 → 输出建议 → 结束。"
    )

    add_heading(doc, "7.8 AI推荐伪代码", 2)
    pseudo = """
输入：user_id
输出：ai_score, model_result, recommendation_text, deepseek_advice

1. 读取用户基本信息、训练记录、饮食记录、身体指标、训练计划；
2. 计算训练频率、饮食记录完整度、身体指标记录完整度、计划完成率、热量平衡；
3. 按权重计算本地AI评分；
4. 根据用户目标判断用户类型：
   如果目标为增肌，则推荐力量训练；
   如果目标为减脂，则推荐有氧和HIIT；
   否则推荐综合塑形方案；
5. 判断DeepSeek API Key是否存在且依赖是否安装；
6. 如果API可用：
      构造Prompt；
      调用DeepSeek接口；
      返回大模型建议；
   否则：
      返回本地推荐结果；
7. 在页面显示评分、用户画像、推荐动作和AI建议。
"""
    add_mermaid_block(doc, "AI推荐伪代码", pseudo)

    # 八 运行设计
    add_heading(doc, "八、运行设计", 1)
    add_heading(doc, "8.1 运行模块组合", 2)
    add_para(doc, "系统启动后首先初始化Flask应用和数据库连接。学生用户通过登录模块进入系统首页，随后可访问训练、饮食、身体指标、计划、评分和AI推荐模块。管理员通过后台登录模块进入后台页面，可进行用户查看、食物库维护和统计信息查看。")

    add_heading(doc, "8.2 运行控制", 2)
    add_table(doc, ["用户状态", "可访问功能", "控制方式"], [
        ["未登录用户", "登录、注册", "若访问业务页面则重定向到登录页。"],
        ["学生用户", "首页、个人信息、训练、饮食、身体指标、计划、评分、AI推荐", "通过session中的user_id判断登录状态。"],
        ["管理员", "后台首页、用户管理、食物库管理", "通过session中的admin_id判断管理员状态。"],
        ["AI接口不可用", "本地评分和本地推荐", "捕获异常并启用降级机制。"]
    ])

    add_heading(doc, "8.3 运行时间要求", 2)
    add_table(doc, ["操作", "期望响应时间", "说明"], [
        ["登录验证", "1秒以内", "本地数据库查询，响应较快。"],
        ["新增记录", "1秒以内", "写入SQLite数据库。"],
        ["数据分析", "2秒以内", "统计用户历史数据并生成图表数据。"],
        ["AI推荐", "3-10秒", "取决于DeepSeek接口响应速度。"],
        ["AI降级推荐", "2秒以内", "无需调用外部接口。"]
    ])

    # 九 出错处理
    add_heading(doc, "九、出错处理与系统维护设计", 1)
    add_heading(doc, "9.1 出错信息设计", 2)
    add_table(doc, ["错误类型", "可能原因", "系统处理方式"], [
        ["登录失败", "用户名或密码错误", "页面提示账号或密码错误。"],
        ["未登录访问", "用户直接访问业务页面", "重定向到登录页面。"],
        ["表单数据错误", "用户输入为空或格式错误", "提示用户重新输入。"],
        ["数据库写入失败", "数据库异常或字段错误", "返回操作失败提示。"],
        ["DeepSeek调用失败", "余额不足、网络异常、API Key错误", "显示错误原因并回退到本地推荐。"],
        ["模板缺失", "HTML文件不存在", "检查templates目录并补充模板文件。"]
    ])

    add_heading(doc, "9.2 补救措施", 2)
    measures = [
        "数据库定期备份，避免误删或损坏造成数据丢失。",
        "关键路由增加登录状态检查，防止未授权访问。",
        "AI调用过程使用try-except捕获异常，避免页面崩溃。",
        "表单输入增加非空校验和默认值处理。",
        "管理员账号与普通学生账号分离，减少越权风险。",
        "后续可对密码进行哈希加密，提高安全性。"
    ]
    for i, item in enumerate(measures, 1):
        add_number(doc, i, item)

    add_heading(doc, "9.3 系统维护设计", 2)
    add_para(doc, "系统采用模块化结构，后续维护可分别从页面层、路由层、模型层和AI服务层进行。若需要扩展新功能，可新增对应数据库表、业务路由和页面模板；若需要替换AI模型，只需修改AI服务调用模块，其他业务模块无需大幅调整。")

    # 十 小结
    add_heading(doc, "十、阶段小结", 1)
    add_para(doc, "通过本概要设计与详细设计阶段，系统已经从需求分析阶段的功能描述进一步转化为可实现的软件结构。本文档完成了数据流图改进、体系架构设计、功能结构图转换、接口设计、数据库表设计和主要算法流程设计，为后续编码实现、系统测试和课程答辩提供了完整依据。")
    add_para(doc, "本系统的设计重点不只是传统信息管理系统的增删改查功能，而是进一步引入数据分析与AI推荐机制，形成“记录—分析—评分—推荐”的闭环，使系统具备更强的现实应用价值和课程设计创新性。")

    # 附录
    add_heading(doc, "附录：主要图表清单", 1)
    add_table(doc, ["图片编号", "图片名称", "建议文件名", "放置位置"], [
        ["图2-1", "改进后的0层数据流图", "图2-1_改进0层数据流图.png", "第二章2.3"],
        ["图2-2", "改进后的1层数据流图", "图2-2_改进1层数据流图.png", "第二章2.4"],
        ["图2-3", "AI推荐子系统数据流图", "图2-3_改进AI推荐数据流图.png", "第二章2.5"],
        ["图2-4", "训练与饮食记录数据流图", "图2-4_改进训练饮食记录数据流图.png", "第二章2.6"],
        ["图3-1", "系统体系架构示意图", "图3-1_系统体系架构示意图.png", "第三章3.2"],
        ["图4-1", "最初的转换结果图", "图4-1_最初转换结果图.png", "第四章4.2"],
        ["图4-2", "最终改进优化后的功能结构图", "图4-2_最终功能结构图.png", "第四章4.3"],
        ["图5-1", "登录界面设计", "图5-1_登录界面设计.png", "第五章5.1"],
        ["图5-2", "系统首页界面设计", "图5-2_系统首页界面设计.png", "第五章5.1"],
        ["图5-3", "训练记录界面设计", "图5-3_训练记录界面设计.png", "第五章5.1"],
        ["图5-4", "饮食记录界面设计", "图5-4_饮食记录界面设计.png", "第五章5.1"],
        ["图5-5", "身体指标界面设计", "图5-5_身体指标界面设计.png", "第五章5.1"],
        ["图5-6", "训练计划界面设计", "图5-6_训练计划界面设计.png", "第五章5.1"],
        ["图5-7", "AI推荐界面设计", "图5-7_AI推荐界面设计.png", "第五章5.1"],
        ["图5-8", "管理员后台界面设计", "图5-8_管理员后台界面设计.png", "第五章5.1"],
        ["图6-1", "系统数据库关系图", "图6-1_系统数据库关系图.png", "第六章6.10"],
        ["图7-1", "用户登录流程图", "图7-1_用户登录流程图.png", "第七章7.1"],
        ["图7-2", "添加训练记录流程图", "图7-2_添加训练记录流程图.png", "第七章7.2"],
        ["图7-3", "添加饮食记录流程图", "图7-3_添加饮食记录流程图.png", "第七章7.3"],
        ["图7-4", "BMI计算流程图", "图7-4_BMI计算流程图.png", "第七章7.4"],
        ["图7-5", "自动生成训练计划流程图", "图7-5_自动生成训练计划流程图.png", "第七章7.5"],
        ["图7-6", "训练质量评分流程图", "图7-6_训练质量评分流程图.png", "第七章7.6"],
        ["图7-7", "AI智能推荐流程图", "图7-7_AI智能推荐流程图.png", "第七章7.7"]
    ])

    doc.save(DOCX_PATH)
    print(f"概要设计与详细设计文档已生成：{DOCX_PATH.resolve()}")
    print(f"图片文件夹路径：{IMAGE_DIR.resolve()}")
    print("请将图片按文档提示命名后放入 images_design 文件夹，再重新运行脚本即可自动插入。")


if __name__ == "__main__":
    build_doc()